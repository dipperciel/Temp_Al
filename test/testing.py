# ---------------------------------------------------------------------------------------------------------- #
# Pre-processes a syllabus for Al the bot
# ---------------------------------------------------------------------------------------------------------- #
import io
import os
import re  # regular expression
from datetime import datetime  # to find the string dates in the nodes and convert them to date objects

import cohere  # to use Cohere, which chooses the best node (from the sorted_nodes_text) for the prompt context
import mammoth  # to convert docx to html
import pandas as pd
# Import packages ------------------------------------------------------------------------------------------ #
from docx import Document  # to import the syllabus (ocx file)
from dotenv import load_dotenv  # install python-dotenv. This is to read .env file containing api keys. Must load python-dotenv
from openai import AzureOpenAI  # to use Azure OpenAI
from openai.types.chat import ChatCompletionUserMessageParam  # to do the chat completions

from conversions import convert_file, find_h_level, find_sections_paragraphs, convert_doc_to_nodes, read_tables, render_tables_add_to_nodes_text, clean_up, convert_to_dict

load_dotenv()  # load .env file

# INPUT variables------------------------------------------------------------------------------------------- #
# json_file_source = "C:/Users/donal/OneDrive - York University/New/Al/Al-E/syllabus.json"
# file_source = "C:/Users/donal/OneDrive - York University/New/Al/Al-E/Syllabus6.docx"
file_source = "C:/Users/donal/OneDrive - York University/New/Roots of Modern Canada/0. General/_FW 2024-2025/Syllabus HUMA 1740 FW (2024-2025).docx"
doc = Document(file_source)
test_queries = "C:/Users/donal/OneDrive - York University/New/Al/Al-E/Queries_Temporal.xlsx"  # contains the test queries, either Queries_Questions.xlsx or Test_Questions.xlsx or Queries_Syllabus.xlsx
testing_results = "C:/Users/donal/OneDrive - York University/New/Al/Al-E/Testing_Results.xlsx"  # output file containing the testing results

# Initial set-up
question_bank = data = pd.read_excel(test_queries, sheet_name='Sheet1')  # grab Excel sheet with questions and create data frame

with open(file_source, "rb") as docx_file:
    result = mammoth.convert_to_html(docx_file)  # this converts my doc to html
    html_text = result.value

# Put the course title and rubric/number in variables for further use
course_number = doc.paragraphs[0].text.strip()
course_title = doc.paragraphs[1].text.strip()


# Functions -------------------------------------------------------------------------------------------------- #


# -------------------------------------------------------------------------------------------------------------#
# Main program ----------------------------------------------------------------------------------------------- #
# -------------------------------------------------------------------------------------------------------------#

def get_test_questions(test_queries):
    question_bank = pd.read_excel(test_queries,
                                  sheet_name='Sheet1')  # grab Excel sheet with questions and create data frame
    test_queries = question_bank['Queries'].tolist()
    return test_queries


def pre_process_query(query, html_text):
    # If the query includes the course number or title, it always selects the Course Information, which is not helpful.
    query = query.replace(course_number, "this course")
    query = query.replace(course_title, "this course")

    # LLM has difficulty ansering questions that start with "I don't know", so it needs to be modified.
    if query.startswith(("I don't know", "I dont know")):
        query = query.replace("I don't know", "I would like to know")
        query = query.replace("I dont know", "I would like to know")

    # add a question mark if the query doesn't end in a question mark and starts with a question word
    question_start = ["what", "how", "why", "when", "who", "whom", "whose", "which", "where", "does", "is", "are",
                      "can", "could", "will", "would", "should", "may", "might", "have", "must"]
    if query[len(query) - 1] != "?":
        if query.split()[0].lower() in question_start:
            query = query.rstrip() + "?"

    # this part is about handling temporal relations with previous, next, first and last
    # assignment_dates = ["October 19, 2023", "November 30, 2023", "February 29, 2024", "April 4, 2024"]
    # html_buffer = StringIO(html_text)  # Have to wrap the HTML string in a StringIO object
    # doc_tables_df = pd.read_html(html_buffer)  # This grabs all the tables in the syllabus and stores them in a dataframe
    # print("doc_tables_df: ", doc_tables_df[evaluation_table_index])
    return query


def launch_cohere(sorted_nodes_text, query):
    co = cohere.Client(os.getenv('COHERE_DEV'))  # prod key in env.env

    # initial value
    docs = sorted_nodes_text

    response = co.rerank(
        model="rerank-english-v3.0",
        query=query,
        documents=docs,
        top_n=3,
    )

    results = str(response)
    pos_index1 = results.find("index=")
    pos_index2 = results.find(",", pos_index1)
    index = int(results[pos_index1 + 6:pos_index2:])
    pos_index1 = results.find("relevance_score=")
    pos_index2 = results.find(")", pos_index1)
    relevance_score = float(results[pos_index1 + 16:pos_index2:])
    prompt_context = sorted_nodes_text[int(index)]

    return relevance_score, prompt_context, index


def find_dates_in_nodes(lines, date_today):  # get the node with the dates and identify the line before today and the line after today's date
    i = 0
    date_dic = {}
    # create a dictionary with the line index and the date (which converted into a date object for calculations)
    for line in lines:
        match = re.search(r"[A-S][a-v]{2} \d{1,2}, 20\d{2}", line)  # matches dates that look like "Nov 4, 2023"
        if match:
            line_idx = i
            date_obj = datetime.strptime(match.group(), '%b %d, %Y')  # convert the date string to a datetime object
            date_dic[line_idx] = date_obj  # add the date to the dictionary
        i = i + 1

    # finds the dictionary key where the date is before today's date
    for key, value in date_dic.items():
        if value <= date_today:
            date_before = key

    # finds the dictionary key where the date is after today's date
    iterator = iter(date_dic)
    for key in iterator:
        if key == date_before:
            break
    date_after = next(iterator, None)
    return date_after, date_before


def add_temporal_relation(query, sorted_nodes_text):  # this function adds temporal relations to the prompt in construct_prompt
    # date_today = datetime.today().strftime('%b %d, %Y')
    date_today = datetime.strptime("Nov 21, 2023", "%b %d, %Y")  # for testing purposes
    temporal_relation = ""
    # find the node for the Summary of Evaluation table and the node for the Schedule and Readings table
    i = 0

    assignment_idx: int = -1
    lecture_idx: int = -1

    for sorted_node in sorted_nodes_text:
        if "*Schedule and Readings*" in sorted_node:
            lecture_idx = i
        elif "*Summary of Evaluation*" in sorted_node:
            assignment_idx = i
        i = i + 1

    # If the query contains the word "assignment", including misspelling (assi\w{1,3}ent)
    if re.search(r"assi\w{1,3}ent", query.lower()) and ("next" in query.lower() or "upcoming" in query.lower()):
        temporal_node = sorted_nodes_text[assignment_idx]  # Uses the node for the Summary of Evaluation table
        lines = temporal_node.splitlines()  # split the node into lines
        date_after, date_before = find_dates_in_nodes(lines, date_today)
        temp_index = lines[int(date_after)].find("is due on")
        temporal_relation = (
                f"Today is {date_today.strftime("%b %d, %Y")} and the next assignment is the" + lines[int(date_after)][4:temp_index] + "and it is due on " + lines[int(date_after)][-13:]
        )

    elif re.search(r"assi\w{1,3}ent", query.lower()) and re.search(r"previo\w*", query.lower()):
        temporal_node = sorted_nodes_text[assignment_idx]  # Uses the node for the Summary of Evaluation table
        lines = temporal_node.splitlines()  # split the node into lines
        date_after, date_before = find_dates_in_nodes(lines, date_today)
        temp_index = lines[int(date_before)].find("is due on")
        temporal_relation = "Today is " + date_today.strftime("%b %d, %Y") + " and the previous assignment was the" + \
                            lines[int(date_before)][4:temp_index] + "and it was due on " + lines[int(date_before)][-13:]

    elif re.search(r"assi\w{1,3}ent", query.lower()) and ("first" in query.lower() or "fisrt" in query.lower()):
        temporal_node = sorted_nodes_text[assignment_idx]  # Uses the node for the Summary of Evaluation table
        lines = temporal_node.splitlines()  # split the node into lines
        temp_index = lines[5].find("is due on")
        temporal_relation = "The first assignment is the" + lines[5][4:temp_index] + "and is due on " + lines[5][
                                                                                                        -13:] + "\n"

    elif re.search(r"assi\w{1,3}ent", query.lower()) and ("last" in query.lower() or "lasst" in query.lower()):
        temporal_node = sorted_nodes_text[assignment_idx]  # Uses the node for the Summary of Evaluation table
        lines = temporal_node.splitlines()  # split the node into lines
        temp_index = lines[-2].find("is due on")
        temporal_relation = "The last assignment is the" + lines[-2][4:temp_index] + "and is due on" + lines[-2][
                                                                                                       -13:] + "\n"

    elif re.search(r"assi\w{1,3}ent", query.lower()) and ("today" in query.lower() or "tody" in query.lower()):
        temporal_node = sorted_nodes_text[assignment_idx]  # Uses the node for the Summary of Evaluation table
        lines = temporal_node.splitlines()  # split the node into lines
        date_after, date_before = find_dates_in_nodes(lines, date_today)
        date_in_line = lines[int(date_before)][-13:-1].strip()
        if datetime.strptime(date_in_line, "%b %d, %Y") == date_today:
            temp_index = lines[int(date_before)].find("is due on")
            temporal_relation = "Today is " + date_today.strftime("%b %d, %Y") + " and today's assignment is the" + \
                                lines[int(date_before)][4:temp_index]
        else:
            temporal_relation = "Today is " + date_today.strftime(
                "%b %d, %Y") + " and there are no assignments due today"

    # If the query is about class content (i.e. the word "assignment" is not in Query)------- #

    else:
        if "next" in query.lower() or "upcoming" in query.lower():
            temporal_node = sorted_nodes_text[lecture_idx]  # Uses the node for the Schedule and Readings table
            lines = temporal_node.splitlines()  # split the node into lines
            date_after, date_before = find_dates_in_nodes(lines, date_today)
            temporal_relation = "Today is " + date_today.strftime("%b %d, %Y") + " and the next topic in class" + lines[
                                                                                                                      int(date_after)][
                                                                                                                  10:] + "\n" + \
                                lines[int(date_after + 1)]

        elif re.search(r"previo\w*", query.lower()):
            temporal_node = sorted_nodes_text[lecture_idx]  # Uses the node for the Schedule and Readings table
            lines = temporal_node.splitlines()  # split the node into lines
            date_after, date_before = find_dates_in_nodes(lines, date_today)
            temporal_relation = "Today is " + date_today.strftime("%b %d, %Y") + " and the previous topic in class" + \
                                lines[int(date_before - 1)][10:] + "\n" + lines[int(date_before)]

        elif "first" in query.lower():
            temporal_node = sorted_nodes_text[lecture_idx]  # Uses the node for the Schedule and Readings table
            lines = temporal_node.splitlines()  # split the node into lines
            temporal_relation = "The first topic presented in class" + lines[1][10:] + "\n" + lines[2]

        elif "last" in query.lower():
            temporal_node = sorted_nodes_text[lecture_idx]  # Uses the node for the Schedule and Readings table
            lines = temporal_node.splitlines()  # split the node into lines
            temporal_relation = "The last topic presented in class" + lines[-2][10:] + "\n" + lines[-3]

        elif "today" in query.lower():
            temporal_node = sorted_nodes_text[lecture_idx]  # Uses the node for the Schedule and Readings table
            lines = temporal_node.splitlines()  # split the node into lines
            date_after, date_before = find_dates_in_nodes(lines, date_today)
            date_in_line = lines[int(date_before - 1)][-13:-1].strip()
            if datetime.strptime(date_in_line, "%b %d, %Y") == date_today:
                temporal_relation = "Today is " + date_today.strftime(
                    "%b %d, %Y") + " and today's topic presented in class" + lines[int(date_before - 1)][10:] + "\n" + \
                                    lines[int(date_before)]
            else:
                temporal_relation = "Today is " + date_today.strftime("%b %d, %Y") + " and there is no class today"
    return temporal_relation


def construct_prompt(query, prompt_context):
    system_prompt = "You are a helpful assistant for this course, " + course_number + " ('" + course_title + "'), at York University."
    header = "Answer the question as truthfully as possible using the provided context. If the answer is not contained within the text below, say \"I don't know.\". If a URL link is in the context, always include it in the response."
    separator = "\n\n=====\n\n"
    # context = prompt_context # actual node provided by Cohere
    if file_source == "C:/Users/donal/OneDrive - York University/New/Al/Al-E/Syllabus3.docx":
        temporal_relation = add_temporal_relation(query, sorted_nodes_text)
    else:
        temporal_relation = ""
    prompt = system_prompt + separator + header + separator + "Context:\n" + prompt_context + "\n\n" + temporal_relation + separator + "Question: " + query + " " + "\n\nAnswer:"
    return prompt


def launch_chat_completion(query, prompt_context):
    client = AzureOpenAI(
        api_key=os.getenv('OPENAI_API_KEY'),
        api_version="2024-02-01",
        azure_endpoint="https://cria-dev-useast.openai.azure.com"
    )
    deployment_name = 'cria-gpt-4o-mini'
    start_phrase = construct_prompt(query, prompt_context)
    response = client.chat.completions.create(
        model=deployment_name,
        messages=[
            ChatCompletionUserMessageParam(
                content=start_phrase,
                role='user'
            )
        ],
        max_tokens=4000,
        temperature=1.0
    )
    completion_response = response.dict()["choices"][0]["message"]["content"]
    return completion_response


if __name__ == '__main__':

    file_bytes: bytes = open(file_source, "rb").read()

    sorted_nodes_text = convert_file(io.BytesIO(file_bytes))

    # This part let's you chose whether to run the program in auto or manual mode
    run_mode = "manual"  # "auto" or "manual". Auto is for auto testing all questions and manual is for individual queries

    if run_mode == "auto":  # start the automatic testing
        test_questions = get_test_questions(test_queries)  # get all the test question from the testing file
        relevance_scores_list = []
        index_list = []
        completion_response_list = []
        for i in range(len(get_test_questions(test_queries))):  # test each question
            query = test_questions[i]
            query = pre_process_query(query, html_text)
            relevance_score, prompt_context, index = launch_cohere(sorted_nodes_text,
                                                                   query)  # get the prompt context for the chat completion
            relevance_scores_list.append(relevance_score)
            index_list.append(index)
            completion_response = launch_chat_completion(query, prompt_context)  # this is where the chat completion happens
            completion_response_list.append(completion_response)  # where the completions are stored

        # Now adding the results to the question_bank dataframe
        question_bank.insert(0, 'Relevance', relevance_scores_list)
        question_bank.insert(1, 'Index', index_list)
        question_bank.insert(2, 'Query', test_questions)
        question_bank.insert(3, 'Answer', completion_response_list)
        # Replace content of testint file with the new dataframe
        question_bank.to_excel(testing_results, sheet_name='Sheet1', index=False, engine='openpyxl')

    else:
        query = "What assignment do we have next?"
        query = pre_process_query(query, html_text)

        relevance_score, prompt_context, index = launch_cohere(sorted_nodes_text,
                                                               query)  # get the prompt context for the chat complettion
        completion_response = launch_chat_completion(query, prompt_context)  # this is where the chat completion happens
        # print("First index: ", index)
        # print("First relevance_score: ", relevance_score)
        print("First query:", query)
        # print("First prompt_context:", prompt_context)
        print("First Answer: ", completion_response)

        # Some cleaning up of the answer
        # if "does not directly address the question" in completion_response:
        #    completion_response = "That is something I can't answer"

        # If the response is "I don't know" from the Questions document, then use Syllabus
        if "I don't know" in completion_response or "does not provide information" in completion_response or relevance_score < 0.01:
            # file_source = "C:/Users/donal/OneDrive - York University/New/Al/Al-E/Syllabus6.docx"
            file_source = "C:/Users/donal/OneDrive - York University/New/Roots of Modern Canada/0. General/_FW 2024-2025/Syllabus HUMA 1740 FW (2024-2025).docx"
            doc = Document(file_source)

            with open(file_source, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)  # this converts my doc to html
                html_text = result.value

            sections = find_h_level(doc)
            section_paragraphs = find_sections_paragraphs(sections, doc)
            nodes_text = convert_doc_to_nodes(section_paragraphs, doc)
            doc_tables_df, table_titles = read_tables(html_text)
            render_tables_add_to_nodes_text(table_titles, nodes_text, doc_tables_df)
            sorted_nodes_text = clean_up(nodes_text, sections)
            json_data = convert_to_dict(sorted_nodes_text, file_source)

            relevance_score, prompt_context, index = launch_cohere(sorted_nodes_text, query)
            completion_response = launch_chat_completion(query, prompt_context)
            # print("Second index: ", index)
            # print("Second relevance_score: ", relevance_score)
            print("Second query:", query)
            # print("Second prompt_context", prompt_context)
            print("Second Answer: ", completion_response)

    # -------------------------------------------------------------------------------------------#
    # Different models for Open
    # -------------------------------------------------------------------------------------------#

    # endpoint = "https://aca-dev.openai.azure.com/openai/deployments/fakesmarts-dev/chat/completions?api-version=2023-05-15"
    # fakesmarts-dev uses model name: Model name: gpt-35-turbo-16k
    # endpoint = "https://aca-dev.openai.azure.com/openai/deployments/fakesmarts-dev-embedding/chat/completions?api-version=2023-05-15"
    # fakesmarts-dev-embedding uses model name: Model name: text-embedding-ada-002
    # endpoint = "https://aca-dev.openai.azure.com/openai/deployments/cria-dev-text-embedding-3-large/chat/completions?api-version=2023-05-15"
    # cria-dev-text-embedding-3-large uses model name: text-embedding-3-large
