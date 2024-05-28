# ---------------------------------------------------------------------------------------------------------- #
# Pre-processes a syllabus for Al the bot
# ---------------------------------------------------------------------------------------------------------- #

# Import packages ------------------------------------------------------------------------------------------ #
from docx import Document  # to import the syllabus (ocx file)
import pandas as pd
import mammoth  # to convert docx to html
from io import StringIO
import re  # regular expression
import json  # to convert a list go json
from bs4 import BeautifulSoup # install beautifulsoup4
import cohere  # to use Cohere, which chooses the best node (from the sorted_nodes_text) for the prompt context
from openai import AzureOpenAI  # to use Azure OpenAI
from openai.types.chat import ChatCompletionUserMessageParam  # to do the chat completions
import openpyxl  # to import excel file containing test questions
from datetime import datetime # to find the string dates in the nodes and convert them to date objects

# INPUT variables------------------------------------------------------------------------------------------- #
# json_file_source = "C:/Users/donal/OneDrive - York University/New/Al/Al-E/syllabus.json"
file_source = "C:/Users/donal/OneDrive - York University/New/Al/Al-E/Questions.docx"
doc = Document(file_source)
test_queries = "C:/Users/donal/OneDrive - York University/New/Al/Al-E/Queries_Temporal.xlsx" # contains the test queries, either Queries_Questions.xlsx or Test_Questions.xlsx or Queries_Syllabus.xlsx
testing_results = "C:/Users/donal/OneDrive - York University/New/Al/Al-E/Testing_Results.xlsx" # output file containing the testing results

# Initial set-up
question_bank = data = pd.read_excel(test_queries, sheet_name='Sheet1')  # grab Excel sheet with questions and create data frame

with open(file_source, "rb") as docx_file:
    result = mammoth.convert_to_html(docx_file)  # this converts my doc to html
    html_text = result.value

# Put the course title and rubric/number in variables for further use
course_number = doc.paragraphs[0].text.strip()
course_title = doc.paragraphs[1].text.strip()

# Functions -------------------------------------------------------------------------------------------------- #


def find_hlevel(doc):
    headings = []
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            headings.append(paragraph.text.strip())
    # Erase empty headers
    sections = [item for item in headings if item]  # iterating item through headings, if item is true (i.e. not empty) add item in the list
    return sections

def find_sections_paragraphs(sections,doc):  # finds the paragraph of the sections defined in the list sections
    section_paragraphs = []
    for i in range(len(sections)):
        for j in range(len(doc.paragraphs)):
            if sections[i] == doc.paragraphs[j].text.strip(): # if the section in the list is the same as the doc paragraph, add the paragraph number to section_paragraph list
                section_paragraphs.append(j)
    return section_paragraphs


def convert_doc_to_nodes(section_paragraphs,doc):
    nodes_text = []
    nodes_temp = ""
    for i in range(len(section_paragraphs)-1):
        for j in range(section_paragraphs[i]+1, section_paragraphs[i+1]):
            # first check if there's a hyperlink in the paragraph
            hyperlink_text, hyperlink_url = include_hyperlink(doc.paragraphs[j])
            for k in range(len(hyperlink_text)): # loop to grab all items in hyperlink_text[] and hyperlink_url[]
                if len(hyperlink_text) > 0:  # if there's no item in hypertext_link or hypertext_url, you get an error message
                    temp_text = doc.paragraphs[j].text.replace(hyperlink_text[k], "["+ hyperlink_text[k] + "](" + hyperlink_url[k] + ")") # follows the mark down format
                    doc.paragraphs[j].text = temp_text
            nodes_temp = nodes_temp + doc.paragraphs[j].text.strip() + " "
        nodes_text.append("*" + sections[i] + "*\n" + nodes_temp + "\n")
        nodes_temp = ""
    # need to add the following code to capture the paragraphs in the last section, which are not captured in the loop because the loop would be out of range
    for i in range(section_paragraphs[len(section_paragraphs)-1], len(doc.paragraphs)):  # from the last section position to the last paragraph of the document
        if i == section_paragraphs[len(section_paragraphs)-1]:
            nodes_temp = nodes_temp + "*" + doc.paragraphs[i].text.strip() + "*\n"
        else:
            # first check if there's a hyperlink in the paragraph
            hyperlink_text, hyperlink_url = include_hyperlink(doc.paragraphs[i])
            for k in range(len(hyperlink_text)):  # loop to grab all items in hyperlink_text[] and hyperlink_url[]
                if len(hyperlink_text) > 0:  # if there's no item in hypertext_link or hypertext_url, you get an error message
                    temp_text = doc.paragraphs[i].text.replace(hyperlink_text[k], "[" + hyperlink_text[k] + "](" + hyperlink_url[k] + ")")  # follows the mark down format
                    doc.paragraphs[i].text = temp_text
            nodes_temp = nodes_temp + doc.paragraphs[i].text.strip()
    nodes_text.append(nodes_temp)
    nodes_temp = ""

    # Add the course title, rubric and number, which are not in Course Information, but in the title
    if "Course Information" in nodes_text[0]:  # The following code does not apply with Questions.docx
        nodes_text[0] = nodes_text[0] + "The course rubric and number is " + doc.paragraphs[0].text.strip() + ".\n"
        nodes_text[0] = nodes_text[0] + "The course title is " + doc.paragraphs[1].text.strip() + "."

    return nodes_text


def read_tables_bs4mp(html_text):
    # Parse the HTML with BeautifulSoup
    soup = BeautifulSoup(html_text, 'html.parser')

    # Find all tables
    tables = soup.find_all('table')

    data_frames = []
    for table in tables:
        # Find all rows
        rows = table.find_all('tr')

        table_data = []
        for row in rows:
            cols = row.find_all('td')

            # Extract text and links
            cols_text = [col.get_text() for col in cols]
            cols_links = [col.find('a')['href'] if col.find('a') and 'href' in col.find('a').attrs else '' for col in cols]

            # Combine text and links
            cols_with_links = [f'[{text}] ({link})' if link else text for text, link in zip(cols_text, cols_links)]

            table_data.append(cols_with_links)

        # Convert table data to DataFrame
        df = pd.DataFrame(table_data)

        data_frames.append(df)

    return data_frames

def read_tables(html_text):
    # With the following two code lines, pd.read_html did not keep the links. So I used beautiful soup instead
    # html_buffer = StringIO(html_text)  # Had to wrap the HTML string in a StringIO object because a direct pd.read (see line below) will be deprecated
    # doc_tables_df = pd.read_html(html_buffer)  # This grabs all the tables in the syllabus and stores them in a dataframe
    doc_tables_df = read_tables_bs4mp(html_text)

    # find all caption
    # pattern = r"</table>(.*?)</p>"  # All captions are right after the table between <p> and </p>
    # matches = re.findall(pattern, html_text)
    # table_titles = []
    # for i in range(len(matches)):
    #     if matches[i][5].isdigit():
    #         table_titles.append(matches[i][5:])  # Create table title without the initial 2 digits
    #     else:
    #         table_titles.append(matches[i][4:])

    # find name of table, which is the heading above it (h1, h2 or h3
    longest_section_len = len(max(sections, key=len)) # first determine the length of the longest section
    pattern_close = r"</h\d>\s*<table>"  # Finds all h1 or h2 or h3 followed by a table
    closing_indexes = [m.start() for m in re.finditer(pattern_close, html_text)]  # Finds the closing index of the table

    previous_index = [] # this goes back to the index at a distance of the longest section
    for closing_index in closing_indexes:
        previous_index.append(closing_index - longest_section_len)
    chunks = []  # chunks contain the titles but addtional crap beforehand
    for i in range(len(closing_indexes)):
        chunks.append(html_text[previous_index[i]:closing_indexes[i]+12])

    almost_titles = []  # an almost title still needs to be cleaned further to get the title
    for item in chunks:
        match = re.search(r"<h\d>.*</h", item)
        if match:
            almost_titles.append(match.group())
    # Clean the almost titles to have real titles
    table_titles = []
    for i in range(len(almost_titles)):  # Removes the initial tag and the ending tag
            table_titles.append(almost_titles[i][4:-3])

    for i in range(len(table_titles)):  # this part is to find occurrence of another heading in the title (one that should be eliminated)
        if "</h" in table_titles[i]:
            temp_index = table_titles[i].find("</h") + 9
            table_titles[i] = table_titles[i][temp_index:]

    return doc_tables_df, table_titles


def render_tables_add_to_nodes_text(table_titles, nodes_text, doc_table_df):
    i = 0
    temp_text = ""
    for title in table_titles:
        temp_df = (doc_tables_df[i])

        if title == "Tutorials":
            temp_text = "*Tutorials*\n "
            temp_text = temp_text + ("Who your TA is and what your TA's email is, and what your tutorial time and day, "
                                     "your tutorial room, and your tutorial Zoom address are depends on which tutorial "
                                     "your are in. If the tutorial information is not provided, please always provide "
                                     "a conditional answer that includes all possibilities. Example of a proper answer: "
                                     "'if you are in Tutorial 1, your TA is...; if you are in Tutorial 2, your TA is...; "
                                     "if you are in Tutorial 3, your TA is...'. \n ")
            for j in range(1, len(temp_df)):
                temp_text = temp_text + "If you are in Tutorial " + temp_df.iloc[j, 0] + ", your TA (or teaching assistant or tutor or responsible instructor who teaches the tutorial) is " + temp_df.iloc[j, 1] + ".\n "
                temp_text = temp_text + "If you are in Tutorial " + temp_df.iloc[j, 0] + ", your tutorial time is " + temp_df.iloc[j, 2] + ".\n "
                temp_text = temp_text + "If you are in Tutorial " + temp_df.iloc[j, 0] + ", your tutorial room is " + temp_df.iloc[j, 3] + ".\n "
                temp_text = temp_text + "If you are in Tutorial " + temp_df.iloc[j, 0] + ", your Zoom address (or Zoom link) during online sessions is " + temp_df.iloc[j, 4] + ".\n"
            nodes_text.append(temp_text)

        elif title == "Faculty Members Information":
            temp_text = "*Faculty Members Information*\n "
            for j in range(1, len(temp_df)):
                temp_text = temp_text + temp_df.iloc[j, 0] + " is the course's " + temp_df.iloc[j, 1] + " and has the following email address: " + temp_df.iloc[j, 2] + " and has the following office hours (time you can meet or appointment time): " + temp_df.iloc[j, 3] + "and has the following office address or location (where you can meet with your professor or instructor or teacher or TA): " + temp_df.iloc[j, 4] + ".\n "
            nodes_text.append(temp_text)

        elif title == "Summary of Evaluation":
            temp_text = "*Summary of Evaluation*\n "
            temp_text = temp_text + "This section answers questions about how much an assignment is worth (how much it counts toward the final grade) and when the assignments are due or have to be submitted or handed in (submission date). \n"
            # evaluation_table_index = i  # this is to be able to find the evaluation table when we do the query precoessing for temporal relations
            for j in range(1, len(temp_df)):
                temp_text = temp_text + "The " + temp_df.iloc[j, 0] + " is worth " + temp_df.iloc[j, 1] + " of the final grade. In other words, it counts for " + temp_df.iloc[j, 1] + " of the final grade.\n "
                temp_text = temp_text + "The " + temp_df.iloc[j, 0] + " is due on " + temp_df.iloc[j, 2] + ". In other words, the deadline or due date or submission date for " + temp_df.iloc[j, 0] + " is " + temp_df.iloc[j, 2] + ".\n "
            nodes_text.append(temp_text)

        elif title == "Grading Equivalence":
            temp_text = "*Grading Equivalence*\n "
            for j in range(1, len(temp_df)):
                temp_text = temp_text + temp_df.iloc[j, 0] + " is the same as a grade point of " + temp_df.iloc[j, 1] + ", which falls in the percent range of " + temp_df.iloc[j, 2] + "%, and is described as '" + temp_df.iloc[j, 3] + "'.\n "
            nodes_text.append(temp_text)

        elif title == "Definitions of Standing":
            temp_text = "*Definitions of Standing*\n "
            for j in range(0, len(temp_df)):
                temp_text = temp_text + "A grade considered '" + temp_df.iloc[j, 0] + "' means that you have a " + temp_df.iloc[j, 1] + "\n "
            nodes_text.append(temp_text)

        elif title == "Schedule and Readings":
            temp_text = "*Schedule and Readings*\n "
            for j in range(1, len(temp_df)):
                temp_text = temp_text + "The topic on " + temp_df.iloc[j, 2] + " is (or is about) '" + temp_df.iloc[j, 0] + "'. In other words, '" + temp_df.iloc[j, 0] + "' is presented in class on " + temp_df.iloc[j, 2] + ".\n "
                if str(temp_df.iloc[j, 1]) == "nan":
                    temp_text = temp_text + "There are no readings on " + temp_df.iloc[j, 2] + ".\n "
                else:
                    temp_text = temp_text + "The reading(s) for the topic called '" + temp_df.iloc[j, 0] + "' on " + \
                                temp_df.iloc[j, 2] + " is (are) the following: " + str(temp_df.iloc[j, 1]) + "\n "
            nodes_text.append(temp_text)

        elif title == "Important Dates":
            temp_text = "*Important Dates*\n "
            for j in range(1, len(temp_df)):
                if "None" in temp_df.iloc[j, 1]:
                    temp_text = temp_text + "There is no " + temp_df.iloc[j, 0] + ".\n "
                else:
                    temp_text = temp_text + temp_df.iloc[j, 0] + " is on " + temp_df.iloc[j, 1] + ".\n "
            nodes_text.append(temp_text)

        else:
            temp_text = "*" + title + "*\n "
            nb_rows = len(temp_df)
            nb_columns = len(temp_df.columns)
            for j in range(1, nb_rows):
                temp_text = temp_text + "The following " + temp_df.iloc[0, 0].lower() + ": " + temp_df.iloc[j, 0] + " has "
                for k in range(1, nb_columns - 1):
                    temp_text = temp_text + "the following " + temp_df.iloc[0, k].lower() + ": " + str(temp_df.iloc[j,k]) + " and has "
                temp_text = temp_text + "the following " + temp_df.iloc[0, k+1].lower() + ": " + str(temp_df.iloc[j, k+1]).strip() + "."
            nodes_text.append(temp_text)
        i = i + 1

    return nodes_text


def clean_up(nodes_text, sections):

    # Render the section Course Information
    temp_text = nodes_text[0].replace("Course Director:", "The course director (or professor or instructor or teacher) for this course is ")
    nodes_text[0] = temp_text
    temp_text = nodes_text[0].replace("Email:", "\n Your course director's email is ")
    nodes_text[0] = temp_text
    temp_text = nodes_text[0].replace("Semester:", "\n The current semester (or term) is ")
    nodes_text[0] = temp_text
    temp_text = nodes_text[0].replace("Lecture time & day:", "\n The lecture (or class) is offered on the following day and time: ")
    nodes_text[0] = temp_text
    temp_text = nodes_text[0].replace("Lecture room:", "\n If you're wondering how to get to your lecture, the lecture (or class) takes place in the following classroom (or location): ")
    nodes_text[0] = temp_text
    temp_text = nodes_text[0].replace("Zoom (Lecture):", "\n Some classes may be offered on Zoom or you may have to attend some classes on Zoom only during unforseen situaitons such as snowstorms or the instructor's illness, in which case the Zoom link (or Zoom address) for the lecture will be ")
    nodes_text[0] = temp_text
    temp_text = nodes_text[0].replace("eClass:", "\n There is an eClass site (the course has been uploaded to eClass) and the eClass link (or address or URL) is ")
    nodes_text[0] = temp_text
    temp_text = nodes_text[0].replace("Office:", "\n What is the course director's (or professor's or instructor's or teacher's) office number (or office address)? Where can I meet him or her? The answer is: ")
    nodes_text[0] = temp_text
    temp_text = nodes_text[0].replace("Office Hours:", "\n The course director's (or professor's or instructor's or teacher's) office hours are ")
    nodes_text[0] = temp_text
    temp_text = nodes_text[0].replace("\t", "")
    nodes_text[0] = temp_text

    #Combine "Tutorials" and "Faculty Members Information" for better results
    tutorials_index = 0
    fac_memberss_index = 0
    for i in range(len(nodes_text)):
        if nodes_text[i].find("*Tutorials*") > -1:
            tutorials_index = i
        if nodes_text[i].find("*Faculty Members Information*") > -1:
            fac_memberss_index = i
    if tutorials_index != 0 and fac_memberss_index != 0:
        nodes_text[tutorials_index] = nodes_text[tutorials_index] + nodes_text[fac_memberss_index]
        del nodes_text[fac_memberss_index]

    # Erase empty nodes
    filtered_nodes_text = [text for text in nodes_text if not (text.endswith("*\n\n") or text.endswith("*\n \n"))]
    nodes_text = filtered_nodes_text

    # Combine two nodes when there is a table and text under the same header
    sorted_nodes_text = sorted(nodes_text)  # If you sort the list, like item will be next to each other
    for i in range(len(nodes_text)-2, -1, -1):
        first_index = sorted_nodes_text[i].find("*")
        second_index = sorted_nodes_text[i].find("*", first_index + 1)  # Starts searching after the first *
        temp_node = sorted_nodes_text[i][:second_index]  # you now have the title of the node
        if temp_node == sorted_nodes_text[i+1][:second_index]:
            sorted_nodes_text[i] = sorted_nodes_text[i] + sorted_nodes_text[i+1][second_index+1:]  # Combine the following item with the previous
            del sorted_nodes_text[i+1]  # And delete the following, now redundant

    # From here on (i.e. after clean_up), we must work with sort_nodes_text instead of nodes_text

    # Erase the caption that appears after the title in sorted_nodes_text
    # i = 0
    # for node in sorted_nodes_text:
    #     first_index = node.find("*")
    #     second_index = node.find("*", first_index + 1)  # Starts searching after the first *
    #     title_length = second_index - first_index - 1
    #     initial_title = node[first_index+1:second_index]
    #     potential_caption_index = node.find(initial_title, second_index, second_index + title_length + 4)
    #     potential_caption = node[potential_caption_index:potential_caption_index + title_length]
    #     if initial_title == potential_caption:
    #         node = "*" + initial_title + "*\n" + node[potential_caption_index + title_length + 1:]
    #     sorted_nodes_text[i] = node
    #     i = i + 1
    return sorted_nodes_text


def include_hyperlink(paragraph):  # This function looks for hyperlinks in a paragraph. If found, returns the list of text that has a link and the list of its url
    # has_hyperlink = False
    hyperlink_text = []
    hyperlink_url = []
    if len(paragraph.hyperlinks) > 0:

        for hyperlink in paragraph.hyperlinks:
            hyperlink_text.append(hyperlink.text)
            hyperlink_url.append(hyperlink.url)

    return hyperlink_text, hyperlink_url


def convert_to_json(sorted_nodes_text, file_source):

    # First remove path from file_source
    filename = file_source.split("/")[-1]
    node_number = 0
    json_nodes_text = []
    for text in sorted_nodes_text:
        node = {
            "node_number": node_number,
            "type": "Narrative text",
            "text": text,
            "metadata": {
                "category_depth": 0,
                "filename": filename,
                "page_number": 1,          # not sure what to do with the page number here
                "languages": ["eng"],
                "filetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
        }
        json_nodes_text.append(node)  # json.dumps converts the list to a JSON string
        node_number = node_number + 1
    return json.dumps(json_nodes_text, indent=4)  # json.dumps converts the list to a JSON string - Add indentation for readability


def launch_cohere(sorted_nodes_text, query):
    # co = cohere.Client('AfR2LZg8hnuAaCxDsbhIK6lgwhEQ8VErlZEJeClI') # dev key
    co = cohere.Client('')  # prod key

    # initial value
    docs = sorted_nodes_text

    response = co.rerank(
        model="rerank-english-v3.0",
        query=query,
        documents=docs,
        top_n=3,
    )

    # print(response)
    results = str(response)
    pos_index1 = results.find("index=")
    pos_index2 = results.find(",", pos_index1)
    index = int(results[pos_index1 + 6:pos_index2:])
    pos_index1 = results.find("relevance_score=")
    pos_index2 = results.find(")", pos_index1)
    relevance_score = float(results[pos_index1 + 16:pos_index2:])

    # print("index: ",  index)
    # print("relevance_score: ", relevance_score)
    prompt_context = sorted_nodes_text[int(index)]

    return relevance_score, prompt_context, index


def launch_chat_completion(query, prompt_context):
    client = AzureOpenAI(
        api_key="a2d6d7e9786f441dbd80942bc9848e7c",
        api_version="2023-05-15",
        azure_endpoint="https://aca-dev.openai.azure.com"
    )
    deployment_name = 'fakesmarts-dev'
    # print('Sending a test completion job')
    start_phrase = construct_prompt(query, prompt_context)
    response = client.chat.completions.create(
        model=deployment_name,
        messages=[
            ChatCompletionUserMessageParam(
                content=start_phrase,
                role='user'
            )
        ],
        max_tokens=1024,
        temperature=1.0
    )
    # print(query)
    completion_response = response.dict()["choices"][0]["message"]["content"]
    return completion_response

# print("Answer: ", completion_response)
# print("index: ",  index)
# print("relevance_score: ", relevance_score)


def find_dates_in_nodes(lines, date_today): # get the node with the dates and identify the line before today and the line after today's date
    i = 0
    date_dic = {}
    # create a dictionary with the line index and the date (which converted into a date object for calculations)
    for line in lines:
        print(i,": ", line)
        match = re.search(r"[A-S][a-v]{2} \d{1,2}, 20\d{2}", line)  # matches dates that look like "Nov 4, 2023"
        if match:
            line_idx = i
            date_obj = datetime.strptime(match.group(), '%b %d, %Y')  # convert the date string to a datetime object
            date_dic[line_idx] = date_obj  # add the date to the dictionary
        i = i + 1

    # finds the dictionary key where the date is before today's date
    for key, value in date_dic.items():
        if value <= date_today:
            date_before = key

    # finds the dictionary key where the date is after today's date
    iterator = iter(date_dic)
    for key in iterator:
        if key == date_before:
            break
    date_after = next(iterator, None)
    return date_after, date_before


def add_temporal_relation(query): # this function adds temporal relations to the prompt in construct_prompt
    # date_today = datetime.today().strftime('%b %d, %Y')
    date_today = datetime.strptime("Nov 21, 2023", "%b %d, %Y") # for testing purposes
    temporal_relation = ""

    # find the node for the Summary of Evaluation table and the node for the Schedule and Readings table
    i = 0
    for sorted_node in sorted_nodes_text:
        if "*Schedule and Readings*" in sorted_node:
            lecture_idx = i
        elif "*Summary of Evaluation*" in sorted_node:
            assignment_idx = i
        i = i + 1

    # If the query contains the word "assignment", including misspelling (assi\w{1,3}ent)
    if re.search(r"assi\w{1,3}ent", query.lower()) and ("next" in query.lower() or "upcoming" in query.lower()):
        temporal_node = sorted_nodes_text[assignment_idx]  # Uses the node for the Summary of Evaluation table
        lines = temporal_node.splitlines()  # split the node into lines
        date_after, date_before = find_dates_in_nodes(lines, date_today)
        temp_index = lines[int(date_after)].find("is due on")
        temporal_relation = "Today is " + date_today.strftime("%b %d, %Y") + " and the next assignment is the" + lines[int(date_after)][4:temp_index] + "and it is due on " + lines[int(date_after)][-13:]

    elif re.search(r"assi\w{1,3}ent", query.lower()) and re.search(r"previo\w*", query.lower()):
        temporal_node = sorted_nodes_text[assignment_idx]  # Uses the node for the Summary of Evaluation table
        lines = temporal_node.splitlines()  # split the node into lines
        date_after, date_before = find_dates_in_nodes(lines, date_today)
        temp_index = lines[int(date_before)].find("is due on")
        temporal_relation = "Today is " + date_today.strftime("%b %d, %Y") + " and the previous assignment was the" + lines[int(date_before)][4:temp_index] + "and it was due on " + lines[int(date_before)][-13:]

    elif re.search(r"assi\w{1,3}ent", query.lower()) and ("first" in query.lower() or "fisrt" in query.lower()):
        temporal_node = sorted_nodes_text[assignment_idx]  # Uses the node for the Summary of Evaluation table
        lines = temporal_node.splitlines()  # split the node into lines
        temp_index = lines[5].find("is due on")
        temporal_relation = "The first assignment is the" + lines[5][4:temp_index] + "and is due on " + lines[5][-13:] + "\n"

    elif re.search(r"assi\w{1,3}ent", query.lower()) and ("last" in query.lower() or "lasst" in query.lower()):
        temporal_node = sorted_nodes_text[assignment_idx]  # Uses the node for the Summary of Evaluation table
        lines = temporal_node.splitlines()  # split the node into lines
        temp_index = lines[-2].find("is due on")
        temporal_relation = "The last assignment is the" + lines[-2][4:temp_index] + "and is due on" + lines[-2][-13:] + "\n"

    elif re.search(r"assi\w{1,3}ent", query.lower()) and ("today" in query.lower() or "tody" in query.lower()):
        temporal_node = sorted_nodes_text[assignment_idx]  # Uses the node for the Summary of Evaluation table
        lines = temporal_node.splitlines()  # split the node into lines
        date_after, date_before = find_dates_in_nodes(lines, date_today)
        date_in_line = lines[int(date_before)][-13:-1].strip()
        if datetime.strptime(date_in_line, "%b %d, %Y") == date_today:
            temp_index = lines[int(date_before)].find("is due on")
            temporal_relation = "Today is " + date_today.strftime("%b %d, %Y") + " and today's assignment is the" + lines[int(date_before)][4:temp_index]
        else:
            temporal_relation = "Today is " + date_today.strftime("%b %d, %Y") + " and there are no assignments due today"

    # If the query is about class content (i.e. the word "assignment" is not in Query)------- #

    else:
        if "next" in query.lower() or "upcoming" in query.lower():
            temporal_node = sorted_nodes_text[lecture_idx] # Uses the node for the Schedule and Readings table
            lines = temporal_node.splitlines()  # split the node into lines
            date_after, date_before = find_dates_in_nodes(lines, date_today)
            temporal_relation = "Today is " + date_today.strftime("%b %d, %Y") + " and the next topic in class" + lines[int(date_after)][10:] + "\n" + lines[int(date_after + 1)]

        elif re.search(r"previo\w*", query.lower()):
            temporal_node = sorted_nodes_text[lecture_idx]  # Uses the node for the Schedule and Readings table
            lines = temporal_node.splitlines()  # split the node into lines
            date_after, date_before = find_dates_in_nodes(lines, date_today)
            temporal_relation = "Today is " + date_today.strftime("%b %d, %Y") + " and the previous topic in class" + lines[int(date_before-1)][10:] + "\n" + lines[int(date_before)]

        elif "first" in query.lower():
            temporal_node = sorted_nodes_text[lecture_idx]  # Uses the node for the Schedule and Readings table
            lines = temporal_node.splitlines()  # split the node into lines
            temporal_relation = "The first topic presented in class" + lines[1][10:] + "\n" + lines[2]

        elif "last" in query.lower():
            temporal_node = sorted_nodes_text[lecture_idx]  # Uses the node for the Schedule and Readings table
            lines = temporal_node.splitlines()  # split the node into lines
            temporal_relation = "The last topic presented in class" + lines[-2][10:] + "\n" + lines[-3]

        elif "today" in query.lower():
            temporal_node = sorted_nodes_text[lecture_idx]  # Uses the node for the Schedule and Readings table
            lines = temporal_node.splitlines()  # split the node into lines
            date_after, date_before = find_dates_in_nodes(lines, date_today)
            date_in_line = lines[int(date_before-1)][-13:-1].strip()
            if datetime.strptime(date_in_line, "%b %d, %Y") == date_today:
                temporal_relation = "Today is " + date_today.strftime("%b %d, %Y") + " and today's topic presented in class" + lines[int(date_before-1)][10:] + "\n" + lines[int(date_before)]
            else:
                temporal_relation = "Today is " + date_today.strftime("%b %d, %Y") + " and there is no class today"
    return temporal_relation

def construct_prompt(query, prompt_context):
    system_prompt = "You are a helpful assistant for this course, " + course_number + " ('" + course_title + "'), at York University."
    header = "Answer the question as truthfully as possible using the provided context. If the answer is not contained within the text below, say \"I don't know.\" If a URL link is in the context, always include it in the response."
    separator = "\n\n=====\n\n"
    context = prompt_context # actual node provided by Cohere

    temporal_relation = add_temporal_relation(query)
    # print("temporal_relation: ", temporal_relation)

    prompt = system_prompt + separator + header + separator + "Context:\n" + prompt_context + "\n\n" + temporal_relation + separator + "Question: " + query + " " + "\n\nAnswer:"
    print('prompt: ', prompt)
    return prompt


def get_test_questions(test_queries):
    question_bank = pd.read_excel(test_queries, sheet_name='Sheet1')  # grab Excel sheet with questions and create data frame
    test_queries = question_bank['Queries'].tolist()
    return test_queries


def pre_process_query(query, html_text):

    # If the query includes the course number or title, it always selects the Course Information, which is not helpful.
    query = query.replace(course_number, "this course")
    query = query.replace(course_title, "this course")

    # LLM has difficulty ansering questions that start with "I don't know", so it needs to be modified.
    if query.startswith(("I don't know", "I dont know")):
        query = query.replace("I don't know", "I would like to know")
        query = query.replace("I dont know", "I would like to know")

    # add a question mark if the query doesn't end in a question mark and starts with a question word
    question_start = ["what", "how", "why", "when", "who", "whom", "whose", "which", "where", "does", "is", "are", "can", "could", "will", "would", "should", "may", "might", "have", "must"]
    if query[len(query) - 1] != "?":
        if query.split()[0].lower() in question_start:
            query = query.rstrip() + "?"

    # this part is about handling temporal relations with previous, next, first and last
    # assignment_dates = ["October 19, 2023", "November 30, 2023", "February 29, 2024", "April 4, 2024"]
    # html_buffer = StringIO(html_text)  # Have to wrap the HTML string in a StringIO object
    # doc_tables_df = pd.read_html(html_buffer)  # This grabs all the tables in the syllabus and stores them in a dataframe
    # print("doc_tables_df: ", doc_tables_df[evaluation_table_index])
    return query

# -------------------------------------------------------------------------------------------------------------#
# Main program ----------------------------------------------------------------------------------------------- #
# -------------------------------------------------------------------------------------------------------------#

sections = find_hlevel(doc)
section_paragraphs = find_sections_paragraphs(sections, doc)  # the sections in the sections list are assigned a paragraph
nodes_text = convert_doc_to_nodes(section_paragraphs, doc)  # the doc is converted to a list of semantic sections containing the text
doc_tables_df, table_titles = read_tables(html_text)# I need the dataframe created in read_table to use in render_tables_add_to_notes, where the dataframe is rendered
# read_tables_temp(html_text)
render_tables_add_to_nodes_text(table_titles, nodes_text, doc_tables_df) # Where the rendering of tables is done and added to the list nodes_text
sorted_nodes_text = clean_up(nodes_text, sections)  # final touches to clean up the list

json_data = convert_to_json(sorted_nodes_text, file_source)
# print(json_data)

# This part let's you chose whether to run the program in auto or manual mode
run_mode = "manual"  # "auto" or "manual". Auto is for auto testing all questions and manual is for individual queries

if run_mode == "auto":  # start the automatic testing
    test_questions = get_test_questions(test_queries)  # get all the test question from the testing file
    relevance_scores_list = []
    index_list = []
    completion_response_list = []
    for i in range(len(get_test_questions(test_queries))):  # test each question
        query = test_questions[i]
        query = pre_process_query(query, html_text)
        relevance_score, prompt_context, index = launch_cohere(sorted_nodes_text, query)  # get the prompt context for the chat completion
        relevance_scores_list.append(relevance_score)
        index_list.append(index)
        completion_response = launch_chat_completion(query, prompt_context)  # this is where the chat completion happens
        completion_response_list.append(completion_response)  # where the completions are stored

    # Now adding the results to the question_bank dataframe
    question_bank.insert(0, 'Relevance', relevance_scores_list)
    question_bank.insert(1, 'Index', index_list)
    question_bank.insert(2, 'Query', test_questions)
    question_bank.insert(3, 'Answer', completion_response_list)
    # Replace content of testint file with the new dataframe
    question_bank.to_excel(testing_results, sheet_name='Sheet1', index=False, engine='openpyxl')

else:
    query =  "What assignment did we have previously?"
    query = pre_process_query(query, html_text)
    print(query)
    relevance_score, prompt_context, index = launch_cohere(sorted_nodes_text, query) # get the prompt context for the chat complettion
    completion_response = launch_chat_completion(query, prompt_context)  # this is where the chat completion happens
    print(index)
    print("Answer: ", completion_response)

    # Some cleaning up of the answer
    if "does not directly address the question" in completion_response:
        completion_response = "That is something I can't answer"

    # If the responsse is "I don't know" from the Questions document, then use Syllabus
    if "I don't know" in completion_response:
        file_source = "C:/Users/donal/OneDrive - York University/New/Al/Al-E/Syllabus3.docx"
        doc = Document(file_source)
        sections = find_hlevel(doc)
        section_paragraphs = find_sections_paragraphs(sections, doc)
        nodes_text = convert_doc_to_nodes(section_paragraphs, doc)
        doc_tables_df, table_titles = read_tables(html_text)
        render_tables_add_to_nodes_text(table_titles, nodes_text, doc_tables_df)
        sorted_nodes_text = clean_up(nodes_text, sections)
        json_data = convert_to_json(sorted_nodes_text, file_source)

        print(query)
        relevance_score, prompt_context, index = launch_cohere(sorted_nodes_text, query)
        print(prompt_context)
        completion_response = launch_chat_completion(query, prompt_context)
        print("Answer: ", completion_response)

# -------------------------------------------------------------------------------------------#
# Different models for Open
# -------------------------------------------------------------------------------------------#

# endpoint = "https://aca-dev.openai.azure.com/openai/deployments/fakesmarts-dev/chat/completions?api-version=2023-05-15"
# fakesmarts-dev uses model name: Model name: gpt-35-turbo-16k
# endpoint = "https://aca-dev.openai.azure.com/openai/deployments/fakesmarts-dev-embedding/chat/completions?api-version=2023-05-15"
# fakesmarts-dev-embedding uses model name: Model name: text-embedding-ada-002
# endpoint = "https://aca-dev.openai.azure.com/openai/deployments/cria-dev-text-embedding-3-large/chat/completions?api-version=2023-05-15"
# cria-dev-text-embedding-3-large uses model name: text-embedding-3-large
